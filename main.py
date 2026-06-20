"""
Sermon Tools — AI-powered sermon research, writing, and evaluation
FastAPI backend using OpenRouter (Claude Sonnet)
"""

import os
import io
import re
import json
import logging
from pathlib import Path
from typing import Optional, List
from datetime import datetime

import httpx
from fastapi import FastAPI, HTTPException, UploadFile, File, Form, Request, Depends
from fastapi.staticfiles import StaticFiles
from fastapi.responses import StreamingResponse, HTMLResponse, Response, RedirectResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv

load_dotenv()

import db
import auth
import mcp_refs
import verses

logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")
logger = logging.getLogger(__name__)

app = FastAPI(title="Sermon Tools")
db.init_db()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

static_path = Path(__file__).parent / "static"
app.mount("/static", StaticFiles(directory=static_path), name="static")

OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY")
MODEL = os.getenv("MODEL", "anthropic/claude-sonnet-4-5")
OPENROUTER_BASE_URL = "https://openrouter.ai/api/v1"

# Admin: emails (comma-separated) that get the admin panel.
ADMIN_EMAILS = {e.strip().lower() for e in os.getenv("ADMIN_EMAILS", "dbarger@imb.org").split(",") if e.strip()}

# Models an admin can assign per user (label → OpenRouter id).
MODEL_CHOICES = [
    {"id": "anthropic/claude-sonnet-4-5", "label": "Claude Sonnet 4.5"},
    {"id": "anthropic/claude-haiku-4-5",  "label": "Claude Haiku 4.5"},
    {"id": "openai/gpt-5-mini",           "label": "GPT-5 mini"},
    {"id": "google/gemini-2.5-flash",     "label": "Gemini 2.5 Flash"},
    {"id": "google/gemini-2.5-pro",       "label": "Gemini 2.5 Pro"},
]
_VALID_MODEL_IDS = {m["id"] for m in MODEL_CHOICES} | {MODEL}

# ─── System Prompts ───────────────────────────────────────────────────────────

RESEARCH_SYSTEM_PROMPT = """You are a biblical research assistant for evangelical Christian pastors. You provide thorough, theologically sound research grounded in Scripture.

You hold to core evangelical convictions:
- Biblical inerrancy and the full authority of Scripture
- Salvation by grace alone, through faith alone, in Christ alone
- The Trinity — Father, Son, Holy Spirit
- The deity, virgin birth, bodily resurrection, and return of Jesus Christ
- The necessity of personal repentance and faith for salvation

When researching passages or topics, provide:
1. Exegetical context — original language insights, key terms, grammatical observations
2. Historical and cultural background — what the original audience would have understood
3. Cross-references and thematic connections across Scripture
4. Main theological themes and doctrinal implications
5. Application angles for a contemporary congregation
6. Common misinterpretations or pitfalls to address
7. A suggested sermon outline structure

Be thorough, scholarly but accessible. Format your response with clear headings using markdown."""

WRITE_SYSTEM_PROMPT = """You are a sermon writing assistant for evangelical Christian pastors. Help pastors craft biblically faithful, clear, and engaging sermons.

You hold to core evangelical convictions:
- Biblical inerrancy and the full authority of Scripture
- Salvation by grace alone, through faith alone, in Christ alone
- The Trinity — Father, Son, Holy Spirit
- The deity, virgin birth, bodily resurrection, and return of Jesus Christ
- The necessity of personal repentance and faith for salvation

When writing a sermon outline and draft:
- Ground every point clearly in Scripture with specific references
- Build toward a single, clear big idea (the sermon's main point)
- Structure: compelling introduction, 2-4 main points, conclusion with a clear call to response
- Include illustration suggestions where helpful (mark as [ILLUSTRATION SUGGESTION])
- Ensure the Gospel is clearly present — who Jesus is, what he did, why it matters
- Write in a pastoral, accessible voice
- Application should be specific, actionable, and drawn from the text
- Format with clear markdown headings for easy reference

Produce a complete outline first, then expand into a full draft."""

EVALUATE_SYSTEM_PROMPT = """You are an evangelical Christian theological reviewer evaluating sermons for biblical groundedness and doctrinal soundness.

Evaluate from an evangelical perspective that holds to:
- Biblical inerrancy and the full authority of Scripture
- Salvation by grace alone, through faith alone, in Christ alone
- The Trinity — Father, Son, Holy Spirit
- The deity, virgin birth, bodily resurrection, and return of Jesus Christ
- The necessity of personal repentance and faith for salvation

Evaluate the sermon across five dimensions. For each, give a score (1–10) and specific observations:

## 1. Biblical Groundedness (1–10)
Are claims supported by Scripture? Are passages used in context, not proof-texted? Is interpretation hermeneutically sound?

## 2. Doctrinal Soundness (1–10)
Is the theology orthodox and evangelical? Are there any errors — explicit or by omission? Is the nature of salvation accurately represented?

## 3. Gospel Clarity (1–10)
Is the Gospel — Christ's death and resurrection for sin — clearly present? Is the call to repentance and faith explicit? Would an unbeliever understand how to respond?

## 4. Application Quality (1–10)
Are applications drawn from the text? Are they specific and actionable? Do they connect Scripture to real life?

## 5. Overall Assessment
- **Strengths**: What is working well and worth keeping
- **Improvements**: Specific, actionable suggestions
- **Theological concerns**: Any doctrinal issues that require correction (be direct but pastoral)
- **Overall Score**: Average of the four dimensions with brief summary

Be honest, constructive, and pastoral. Your goal is to help the pastor grow."""


def _lang_directive(lang: Optional[str]) -> str:
    """Extra system-prompt instruction so AI output is generated in the requested
    language. English (the default) adds nothing; Spanish forces a full Spanish
    response. Appended to whichever system prompt an endpoint is using."""
    if (lang or "en").lower().startswith("es"):
        return (
            "\n\nIMPORTANT: Respond entirely in Spanish (neutral Latin American Spanish "
            "suitable for a congregation). All headings, prose, terminology, and analysis "
            "must be in Spanish. When quoting Scripture, use the pastor's preferred Spanish "
            "version."
        )
    return ""


def _passage_block(passage_text: Optional[str], attribution: Optional[str]) -> str:
    """Prompt block carrying the verbatim passage text (when the frontend fetched
    it). Grounds the model in the exact wording instead of its memory. Empty when
    no verbatim text was available, so behavior is unchanged in that case."""
    text = (passage_text or "").strip()
    if not text:
        return ""
    block = (
        "\n\nVERBATIM PASSAGE TEXT (use this exact wording whenever you quote the passage; "
        "do not paraphrase it):\n"
        f"{text}"
    )
    if attribution and attribution.strip():
        block += f"\n[{attribution.strip()}]"
    return block


# ─── File Parsers ─────────────────────────────────────────────────────────────

def extract_text_from_docx(file_bytes: bytes) -> str:
    try:
        import docx
        doc = docx.Document(io.BytesIO(file_bytes))
        return "\n".join(para.text for para in doc.paragraphs if para.text.strip())
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read DOCX file: {e}")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    try:
        import fitz
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        return "".join(page.get_text() for page in doc)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Could not read PDF file: {e}")


# ─── Admin / caller resolution ──────────────────────────────────────────────

def is_admin(user: dict) -> bool:
    return bool(user and (user.get("email") or "").lower() in ADMIN_EMAILS)


def require_admin(http: Request) -> dict:
    uid = auth.get_optional_user_id(http)
    user = db.get_user(uid) if uid else None
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    if not is_admin(user):
        raise HTTPException(status_code=403, detail="Admin access required")
    return user


def _resolve_caller(http: Request):
    """For generating endpoints: require a signed-in user. Returns (user_id,
    model). Raises 401 if not signed in (the tool is sign-in-only), 403 if the
    user is blocked. Per-user assigned model wins over the global default."""
    uid = auth.get_optional_user_id(http)
    user = db.get_user(uid) if uid else None
    if not user:
        raise HTTPException(status_code=401, detail="Please sign in to use Sermon Tools.")
    if user.get("blocked"):
        raise HTTPException(status_code=403, detail="Your access has been suspended. Please contact the administrator.")
    return uid, (user.get("assigned_model") or MODEL)


# ─── Usage logging ────────────────────────────────────────────────────────────

def _record_usage(user_id, feature, model, usage):
    """Persist token/cost accounting from an OpenRouter response. Best-effort —
    never let a logging failure affect the user's request."""
    if not usage:
        return
    try:
        db.log_usage(
            user_id, feature, model,
            usage.get("prompt_tokens"), usage.get("completion_tokens"),
            usage.get("total_tokens"), usage.get("cost"),
        )
    except Exception as e:
        logger.warning("usage log failed: %s", e)


# ─── Streaming Helper ─────────────────────────────────────────────────────────

async def stream_openrouter(
    system_prompt: str, user_message: str, max_tokens: int = 8192,
    model: str = None, user_id: int = None, feature: str = None,
) -> StreamingResponse:
    if not OPENROUTER_API_KEY:
        raise HTTPException(status_code=500, detail="OPENROUTER_API_KEY is not configured")

    use_model = model or MODEL
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://sermon-tools.app",
        "X-Title": "Sermon Tools",
    }

    payload = {
        "model": use_model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ],
        "stream": True,
        "max_tokens": max_tokens,
        # Ask OpenRouter to include token/cost accounting in the final chunk.
        "usage": {"include": True},
    }

    async def generate():
        # Short connect/write timeouts, generous read: once tokens start flowing,
        # the read timeout is the max gap *between* chunks, not a total cap, so a
        # long generation won't be killed mid-stream.
        timeout = httpx.Timeout(connect=15.0, read=300.0, write=15.0, pool=15.0)
        usage = None
        try:
            async with httpx.AsyncClient(timeout=timeout) as client:
                async with client.stream(
                    "POST",
                    f"{OPENROUTER_BASE_URL}/chat/completions",
                    headers=headers,
                    json=payload,
                ) as response:
                    if response.status_code != 200:
                        error_body = await response.aread()
                        logger.error(f"OpenRouter error {response.status_code}: {error_body}")
                        yield f"data: {json.dumps({'error': f'AI service error: {response.status_code}'})}\n\n"
                        return

                    finish_reason = None
                    async for line in response.aiter_lines():
                        if not line.startswith("data: "):
                            continue
                        data = line[6:].strip()
                        if data == "[DONE]":
                            break
                        try:
                            chunk = json.loads(data)
                            if chunk.get("usage"):
                                usage = chunk["usage"]
                            choices = chunk.get("choices") or []
                            if not choices:
                                continue
                            choice = choices[0]
                            if choice.get("finish_reason"):
                                finish_reason = choice["finish_reason"]
                            delta = choice["delta"].get("content", "")
                            if delta:
                                yield f"data: {json.dumps({'text': delta})}\n\n"
                        except (json.JSONDecodeError, KeyError):
                            continue

                    # Don't truncate silently: tell the user (and the log) when the
                    # model hit the output-length ceiling so they can regenerate.
                    if finish_reason == "length":
                        logger.warning("OpenRouter output truncated (finish_reason=length, max_tokens=%s)", max_tokens)
                        note = "\n\n_(This response hit the length limit and was cut off. Use Regenerate, or switch the brief to Concise, to get the rest.)_"
                        yield f"data: {json.dumps({'text': note})}\n\n"
                    yield "data: [DONE]\n\n"
        except Exception as e:
            logger.error(f"Streaming error: {e}")
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
        finally:
            _record_usage(user_id, feature, use_model, usage)

    return StreamingResponse(
        generate(),
        media_type="text/event-stream",
        # Defeat proxy buffering so chunks reach the browser as they're produced.
        headers={"X-Accel-Buffering": "no", "Cache-Control": "no-cache"},
    )


# ─── Routes ───────────────────────────────────────────────────────────────────

@app.get("/", response_class=HTMLResponse)
async def serve_index():
    return (static_path / "index.html").read_text()


class ResearchRequest(BaseModel):
    passage: str
    topic: Optional[str] = None
    notes: Optional[str] = None


class WriteRequest(BaseModel):
    passage: str
    title: Optional[str] = None
    audience: Optional[str] = None
    research_notes: Optional[str] = None
    sermon_length: str = "30-40 minutes"
    style: Optional[str] = None
    lang: Optional[str] = "en"
    passage_text: Optional[str] = None
    passage_attribution: Optional[str] = None


@app.post("/api/research")
async def research(request: ResearchRequest, http: Request):
    if not request.passage.strip():
        raise HTTPException(status_code=400, detail="Scripture passage is required")

    user_message = f"Please research the following for a sermon:\n\nPassage: {request.passage}"
    if request.topic:
        user_message += f"\nTopic focus: {request.topic}"
    if request.notes:
        user_message += f"\nAdditional context: {request.notes}"

    uid, model = _resolve_caller(http)
    return await stream_openrouter(RESEARCH_SYSTEM_PROMPT, user_message, model=model, user_id=uid, feature="research")


@app.post("/api/write")
async def write_sermon(request: WriteRequest, http: Request):
    if not request.passage.strip():
        raise HTTPException(status_code=400, detail="Scripture passage is required")

    user_message = f"Help me write a sermon outline and draft.\n\nMain Passage: {request.passage}"
    if request.title:
        user_message += f"\nSermon Title: {request.title}"
    if request.audience:
        user_message += f"\nCongregation context: {request.audience}"
    if request.research_notes:
        user_message += f"\nResearch notes to incorporate:\n\n{request.research_notes}"
    user_message += f"\nTarget length: {request.sermon_length}"

    if request.style and request.style in SERMON_STYLES:
        s = SERMON_STYLES[request.style]
        user_message += f"""

---
PREACHING STYLE SELECTED: {s['name']}

STRUCTURE:
{s['structure']}

STYLE & VOICE:
{s['style']}

DELIVERY GOAL:
{s['delivery']}

Build this sermon strictly according to these structural and stylistic guidelines. Let the selected approach shape everything — the outline, the pacing, the language, and how truth is revealed."""

    system = WRITE_SYSTEM_PROMPT + _lang_directive(request.lang)
    uid, model = _resolve_caller(http)
    return await stream_openrouter(system, user_message, max_tokens=16000, model=model, user_id=uid, feature="write")


# ─── Section-by-section sermon writing ──────────────────────────────────────────
# Instead of generating a whole 30-40 min sermon in one call (slow, and bumps the
# output ceiling), we generate the outline first, then stream each section in its
# own call. No single call has to fit the whole sermon, so nothing truncates.

def _write_context(req) -> str:
    ctx = f"Main Passage: {req.passage}"
    if req.title:
        ctx += f"\nSermon Title: {req.title}"
    if req.audience:
        ctx += f"\nCongregation context: {req.audience}"
    if req.research_notes:
        ctx += f"\nResearch notes to incorporate:\n{req.research_notes}"
    ctx += f"\nTarget length: {req.sermon_length}"
    if req.style and req.style in SERMON_STYLES:
        s = SERMON_STYLES[req.style]
        ctx += (
            f"\n\nPREACHING STYLE SELECTED: {s['name']}"
            f"\nSTRUCTURE:\n{s['structure']}"
            f"\nSTYLE & VOICE:\n{s['style']}"
            f"\nDELIVERY GOAL:\n{s['delivery']}"
        )
    ctx += _passage_block(getattr(req, "passage_text", None), getattr(req, "passage_attribution", None))
    return ctx


async def complete_openrouter(system_prompt: str, user_message: str, max_tokens: int = 1024,
                              model: str = None, user_id: int = None, feature: str = None) -> str:
    if not OPENROUTER_API_KEY:
        raise HTTPException(status_code=500, detail="OPENROUTER_API_KEY is not configured")
    use_model = model or MODEL
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": "https://sermon-tools.app",
        "X-Title": "Sermon Tools",
    }
    payload = {
        "model": use_model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_message},
        ],
        "max_tokens": max_tokens,
        "usage": {"include": True},
    }
    timeout = httpx.Timeout(connect=15.0, read=120.0, write=15.0, pool=15.0)
    async with httpx.AsyncClient(timeout=timeout) as client:
        resp = await client.post(
            f"{OPENROUTER_BASE_URL}/chat/completions", headers=headers, json=payload
        )
        if resp.status_code != 200:
            logger.error(f"OpenRouter error {resp.status_code}: {resp.text[:500]}")
            raise HTTPException(status_code=502, detail=f"AI service error: {resp.status_code}")
        data = resp.json()
        _record_usage(user_id, feature, use_model, data.get("usage"))
        return data["choices"][0]["message"]["content"]


WRITE_OUTLINE_SYSTEM = WRITE_SYSTEM_PROMPT + """

FOR THIS TASK: produce ONLY the sermon's section outline, as a JSON array of section headings in order — nothing else, no prose, no code fence. Include a compelling introduction, 2-4 main points (each heading should state the actual point, not just "Point 1"), and a conclusion with a clear call to response. Example:
["Introduction: The Ache Nothing Seems to Fill", "Point 1: ...", "Point 2: ...", "Conclusion: ..."]"""


WRITE_SECTION_SYSTEM = WRITE_SYSTEM_PROMPT + """

FOR THIS TASK: write ONLY the single section requested, in full preachable prose (not bullet notes), flowing naturally from the sections already written. Do NOT write any other section, and do not repeat content already covered. Begin with the section heading as a markdown H2 (## ), then the prose. Keep Scripture references specific, the Gospel clear, and application concrete.

LENGTH: size this section for its place in the whole. Across all the sections, the finished sermon should land near the requested target length, so a single section is only a fraction of that, not a whole sermon. Be substantive without padding; let short sections (intros, transitions) stay short."""


def _parse_outline(raw: str) -> List[str]:
    text = raw.strip()
    if text.startswith("```"):
        text = re.sub(r"^```[a-zA-Z]*\n?", "", text)
        text = re.sub(r"\n?```$", "", text).strip()
    start, end = text.find("["), text.rfind("]")
    if start != -1 and end != -1 and end > start:
        text = text[start:end + 1]
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        return []
    return [str(s).strip() for s in data if isinstance(s, (str, int)) and str(s).strip()]


@app.post("/api/write/outline")
async def write_outline(request: WriteRequest, http: Request):
    if not request.passage.strip():
        raise HTTPException(status_code=400, detail="Scripture passage is required")
    uid, model = _resolve_caller(http)
    system = WRITE_OUTLINE_SYSTEM + _lang_directive(request.lang)
    raw = await complete_openrouter(system, _write_context(request), max_tokens=1024,
                                    model=model, user_id=uid, feature="write_outline")
    sections = _parse_outline(raw)
    if not sections:
        raise HTTPException(status_code=502, detail="Could not generate an outline. Please try again.")
    return {"sections": sections}


class WriteSectionRequest(BaseModel):
    passage: str
    title: Optional[str] = None
    audience: Optional[str] = None
    research_notes: Optional[str] = None
    sermon_length: str = "30-40 minutes"
    style: Optional[str] = None
    outline: List[str] = []
    section_title: str
    prior_sections: Optional[str] = None
    lang: Optional[str] = "en"
    passage_text: Optional[str] = None
    passage_attribution: Optional[str] = None


@app.post("/api/write/section")
async def write_section(request: WriteSectionRequest, http: Request):
    if not request.passage.strip():
        raise HTTPException(status_code=400, detail="Scripture passage is required")
    if not request.section_title.strip():
        raise HTTPException(status_code=400, detail="A section title is required")

    msg = _write_context(request)
    if request.outline:
        msg += "\n\nFull sermon outline:\n" + "\n".join(f"- {h}" for h in request.outline)
    if request.prior_sections:
        msg += f"\n\nSections written so far (for continuity — do not repeat these):\n{request.prior_sections}"
    msg += f"\n\nNow write this one section in full:\n{request.section_title}"

    system = WRITE_SECTION_SYSTEM + _lang_directive(request.lang)
    uid, model = _resolve_caller(http)
    return await stream_openrouter(system, msg, max_tokens=8192, model=model, user_id=uid, feature="write_section")


@app.post("/api/evaluate")
async def evaluate_sermon(
    http: Request,
    sermon_text: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    lang: Optional[str] = Form("en"),
):
    text = ""

    if file and file.filename:
        file_bytes = await file.read()
        fname = file.filename.lower()
        if fname.endswith(".docx"):
            text = extract_text_from_docx(file_bytes)
        elif fname.endswith(".pdf"):
            text = extract_text_from_pdf(file_bytes)
        elif fname.endswith(".txt"):
            text = file_bytes.decode("utf-8", errors="replace")
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Use .txt, .docx, or .pdf")
    elif sermon_text:
        text = sermon_text
    else:
        raise HTTPException(status_code=400, detail="Provide sermon text or upload a file")

    if len(text.strip()) < 100:
        raise HTTPException(status_code=400, detail="Sermon text is too short to evaluate meaningfully")

    system = EVALUATE_SYSTEM_PROMPT + _lang_directive(lang)
    uid, model = _resolve_caller(http)
    return await stream_openrouter(system, f"Please evaluate this sermon:\n\n{text}",
                                   model=model, user_id=uid, feature="evaluate")


# ─── Research Steps ──────────────────────────────────────────────────────────

RESEARCH_STEP_BASE = """You are a biblical research assistant for evangelical Christian pastors. You hold to core evangelical convictions: biblical inerrancy, salvation by grace alone through faith alone in Christ alone, the Trinity, the deity and resurrection of Jesus Christ, and the necessity of personal repentance and faith.

Be thorough, scholarly but accessible. Format with clear markdown headings. Focus only on the step described below — do not try to cover other areas."""

RESEARCH_STEPS = [
    {
        "number": 1,
        "title": "Historical & Cultural Context",
        "instruction": "Research the Historical & Cultural Context of this passage only. Cover: Who wrote it, to whom, when, and why. What historical events shaped it. What the original audience would have understood. Relevant geography, politics, Jewish customs, or Greco-Roman context. Keep your focus here — not on word study or application.",
    },
    {
        "number": 2,
        "title": "Text Exegesis",
        "instruction": "Perform Text Exegesis on this passage only. Cover: Verse-by-verse analysis. Key Hebrew/Greek terms and their meaning. Grammatical observations that affect interpretation. How the text's structure shapes its meaning. Any significant textual notes. Keep your focus here — not on background or application.",
    },
    {
        "number": 3,
        "title": "Cross-References & Biblical Connections",
        "instruction": "Research Cross-References and Biblical Connections for this passage only. Cover: Parallel passages in Scripture. Old Testament connections or allusions. How this text fits the biblical-theological storyline (Creation → Fall → Redemption → New Creation). Key canonical themes it develops. Keep your focus here — not on history or application.",
    },
    {
        "number": 4,
        "title": "Theological Themes",
        "instruction": "Identify the Theological Themes of this passage only. Cover: Major doctrinal threads (God's character, humanity, salvation, Christology, eschatology, etc.). How this text contributes to systematic theology. Any tensions or nuances that need careful handling. Heresies or errors this passage addresses or guards against. Keep your focus here — not on application or preaching.",
    },
    {
        "number": 5,
        "title": "Application & Sermon Angles",
        "instruction": "Develop Application and Sermon Angles from this passage only. Cover: How this text speaks to a contemporary congregation. Specific application points (transformational, behavioral, attitudinal). 2-3 different sermon angle options (narrative, topical, expository). Illustration angles or contemporary connections. A suggested sermon outline structure. Keep your focus here — practical and preachable.",
    },
]

STEP_MAP = {s["number"]: s for s in RESEARCH_STEPS}


class ResearchStepRequest(BaseModel):
    step: int
    passage: str
    topic: Optional[str] = None
    notes: Optional[str] = None
    prior_steps: Optional[str] = None
    translation: Optional[str] = None
    brief_type: Optional[str] = None
    lang: Optional[str] = "en"
    passage_text: Optional[str] = None
    passage_attribution: Optional[str] = None


@app.post("/api/research/step")
async def research_step(request: ResearchStepRequest, http: Request):
    if not request.passage.strip():
        raise HTTPException(status_code=400, detail="Scripture passage is required")
    step_def = STEP_MAP.get(request.step)
    if not step_def:
        raise HTTPException(status_code=400, detail="Step must be 1–5")

    system = f"{RESEARCH_STEP_BASE}\n\nYOUR TASK FOR THIS STEP:\n{step_def['instruction']}"

    if request.translation:
        system += f"\n\nThe pastor prefers the {request.translation} translation. Reference this translation when quoting Scripture in your response."
    if request.brief_type == "concise":
        system += "\n\nLength guidance: Be concise — limit this section to 3–4 paragraphs. Prioritize the most exegetically and pastorally significant observations."
    else:
        system += "\n\nLength guidance: Be thorough but disciplined. Develop the most significant observations in real depth, but don't pad, repeat, or pile on exhaustive lists. Keep this section to roughly 800–1,200 words."

    # Ground steps 1–4 in real scholarship: fetch authoritative references from the
    # Study Bible MCP and inject them so the model cites real sources instead of
    # paraphrasing from training data. Returns "" on miss/timeout/disabled (graceful).
    refs = await mcp_refs.fetch_step_refs(request.step, request.passage)
    if refs:
        system += f"\n\n{refs}"

    system += _lang_directive(request.lang)
    system += _passage_block(request.passage_text, request.passage_attribution)

    msg = f"Passage: {request.passage}"
    if request.topic:
        msg += f"\nTopic focus: {request.topic}"
    if request.notes:
        msg += f"\nAdditional context: {request.notes}"
    if request.prior_steps:
        msg += f"\n\nFor context, here is what has been researched in prior steps:\n{request.prior_steps}"
    msg += f"\n\nNow complete Step {request.step}: {step_def['title']}."

    # 16k so an expanded-brief step (the longest single output in the app) doesn't truncate.
    uid, model = _resolve_caller(http)
    return await stream_openrouter(system, msg, max_tokens=16000, model=model, user_id=uid, feature="research_step")


# ─── Verbatim Scripture text ────────────────────────────────────────────────────

@app.get("/api/passage")
async def get_passage(ref: str, translation: str):
    """Verbatim passage text for the chosen translation, when a configured
    provider can serve it. Returns {} (HTTP 200) on a miss so the frontend
    silently falls back to model-quoted Scripture."""
    result = await verses.fetch_passage(ref, translation)
    return result or {}


@app.get("/api/verse-translations")
async def verse_translations():
    """Translation codes that can be served verbatim given the configured keys —
    lets the UI mark those options. Grows automatically if more keys are added."""
    return {"verbatim": sorted(verses.verbatim_capable())}


# ─── Auth Routes ──────────────────────────────────────────────────────────────

@app.get("/api/auth/google")
async def google_login(response: Response):
    url, state = auth.build_authorize_url()
    redirect = RedirectResponse(url=url)
    redirect.set_cookie(
        auth.OAUTH_STATE_COOKIE, state,
        max_age=600, httponly=True, secure=auth.COOKIE_SECURE, samesite="lax"
    )
    return redirect


@app.get("/api/auth/google/callback")
async def google_callback(request: Request, code: str = "", state: str = "", error: str = ""):
    if error:
        return RedirectResponse("/?auth=error")
    saved_state = request.cookies.get(auth.OAUTH_STATE_COOKIE)
    if not saved_state or saved_state != state:
        return RedirectResponse("/?auth=error")
    try:
        userinfo = await auth.exchange_code_for_userinfo(code)
    except Exception as e:
        logger.error(f"OAuth exchange failed: {e}")
        return RedirectResponse("/?auth=error")
    user = db.get_or_create_user(
        google_id=userinfo["sub"],
        email=userinfo.get("email", ""),
        name=userinfo.get("name", ""),
        picture=userinfo.get("picture", ""),
    )
    token = auth.issue_session(user["id"])
    redirect = RedirectResponse("/?auth=ok")
    redirect.set_cookie(
        auth.SESSION_COOKIE, token,
        max_age=auth.COOKIE_MAX_AGE, httponly=True,
        secure=auth.COOKIE_SECURE, samesite="lax"
    )
    redirect.delete_cookie(auth.OAUTH_STATE_COOKIE)
    return redirect


@app.get("/api/auth/me")
async def get_me(request: Request):
    uid = auth.get_optional_user_id(request)
    if not uid:
        return {"user": None}
    user = db.get_user(uid)
    if not user:
        return {"user": None}
    return {"user": {"id": user["id"], "name": user["name"], "email": user["email"],
                     "picture": user["picture"], "is_admin": is_admin(user)}}


@app.post("/api/auth/logout")
async def logout(response: Response):
    response.delete_cookie(auth.SESSION_COOKIE)
    return {"ok": True}


# ─── User Settings ────────────────────────────────────────────────────────────

class SettingsRequest(BaseModel):
    lang: Optional[str] = None
    translation: Optional[str] = None
    length: Optional[str] = None
    style: Optional[str] = None


def _settings_payload(user: dict) -> dict:
    return {
        "lang": user.get("pref_lang"),
        "translation": user.get("pref_translation"),
        "length": user.get("pref_length"),
        "style": user.get("pref_style"),
    }


@app.get("/api/settings")
async def get_settings(request: Request):
    uid = auth.require_user_id(request)
    user = db.get_user(uid)
    if not user:
        raise HTTPException(404, "User not found")
    return _settings_payload(user)


@app.put("/api/settings")
async def put_settings(request: Request, body: SettingsRequest):
    uid = auth.require_user_id(request)
    prefs = {}
    if body.lang is not None:
        prefs["pref_lang"] = "es" if str(body.lang).lower().startswith("es") else "en"
    if body.translation is not None:
        prefs["pref_translation"] = body.translation.strip()[:20] or None
    if body.length is not None:
        prefs["pref_length"] = body.length.strip()[:40] or None
    if body.style is not None:
        prefs["pref_style"] = body.style.strip()[:60] or None
    db.update_user_prefs(uid, **prefs)
    user = db.get_user(uid)
    return _settings_payload(user)


# ─── Admin ────────────────────────────────────────────────────────────────────

class BlockRequest(BaseModel):
    blocked: bool


class ModelRequest(BaseModel):
    model: Optional[str] = None


@app.get("/api/admin/models")
async def admin_models(request: Request):
    require_admin(request)
    return {"models": MODEL_CHOICES, "default": MODEL}


@app.get("/api/admin/users")
async def admin_users(request: Request):
    require_admin(request)
    return db.list_users_admin()


@app.post("/api/admin/users/{user_id}/block")
async def admin_block(request: Request, user_id: int, body: BlockRequest):
    require_admin(request)
    db.set_blocked(user_id, body.blocked)
    return {"ok": True, "blocked": body.blocked}


@app.post("/api/admin/users/{user_id}/model")
async def admin_set_model(request: Request, user_id: int, body: ModelRequest):
    require_admin(request)
    model = body.model or None
    if model and model not in _VALID_MODEL_IDS:
        raise HTTPException(status_code=400, detail="Unknown model")
    db.set_assigned_model(user_id, model)
    return {"ok": True, "model": model}


@app.get("/api/admin/users/{user_id}/sermons")
async def admin_user_sermons(request: Request, user_id: int):
    require_admin(request)
    return db.list_sermons(user_id)


# ─── Sermon Save/Load Routes ──────────────────────────────────────────────────

class SaveSermonRequest(BaseModel):
    title: Optional[str] = None
    passage: str
    topic: Optional[str] = None
    research: Optional[str] = None
    steps: Optional[list] = None


@app.get("/api/sermons")
async def list_sermons(request: Request):
    uid = auth.require_user_id(request)
    return db.list_sermons(uid)


@app.post("/api/sermons")
async def save_sermon(request: Request, body: SaveSermonRequest):
    uid = auth.require_user_id(request)
    title = body.title or (f"{body.passage} Research" if body.passage else "Untitled")
    sermon_id = db.save_sermon(
        user_id=uid,
        title=title,
        passage=body.passage or "",
        topic=body.topic or "",
        research=body.research or "",
        steps=body.steps or [],
    )
    return {"id": sermon_id, "ok": True}


@app.get("/api/sermons/{sermon_id}")
async def get_sermon(request: Request, sermon_id: int):
    uid = auth.require_user_id(request)
    sermon = db.get_sermon(sermon_id, uid)
    if not sermon:
        raise HTTPException(404, "Sermon not found")
    return sermon


# ─── Sermon Styles Config ─────────────────────────────────────────────────────

SERMON_STYLES = {
    "expository-hb-charles": {
        "name": "Expository (Structured & Text-Driven)",
        "structure": "Text-driven, verse-by-verse. Each main point emerges directly from the passage. The text sets the agenda — not the preacher.",
        "style": "Structured and disciplined with crisp, clear transitions. Concise phrasing. Precision in handling the text. No excess — every word serves the passage.",
        "delivery": "The congregation should leave with a clear, ordered understanding of what the text says, what it means, and what it requires of them.",
    },
    "big-idea-tony-evans": {
        "name": "Big Idea (Clear & Authoritative)",
        "structure": "A single dominant idea governs everything. Every point, illustration, and application flows from and returns to that one big idea.",
        "style": "Authoritative and theologically grounded. Culturally engaged — connects biblical truth directly to contemporary life. Confident, declarative tone.",
        "delivery": "The congregation should leave able to state the sermon's one big idea in a single sentence — and feel the weight of it.",
    },
    "narrative-terry-anderson": {
        "name": "Narrative (Story-Driven & Dynamic)",
        "structure": "A full narrative arc: setup → rising tension → climax → resolution. Let the story unfold progressively. Delay the full meaning.",
        "style": "Dynamic pacing with rhythmic phrasing. Build emotional intensity across the sermon. Vivid language and strong sensory detail. Celebratory climax.",
        "delivery": "The congregation should feel like they lived inside the story — not just heard it explained.",
    },
    "inductive-robert-smith": {
        "name": "Inductive (Experiential Discovery)",
        "structure": "Open with tension or a probing question. Withhold resolution. Let the text lead the congregation to discover truth rather than announcing it upfront.",
        "style": "Experiential and imaginative. Embodied storytelling — the preacher inhabits the text. Rich with metaphor, imagery, and incarnational language.",
        "delivery": "The congregation should feel like they discovered the truth themselves — which makes it impossible to forget.",
    },
    "expository-ralph-west": {
        "name": "Expository (Deep & Pastoral)",
        "structure": "Text-driven with layered theological insight. Points unfold with pastoral patience. Depth without rushing. Space for the congregation to sit in the text.",
        "style": "Pastoral, rich, and reflective. Shepherding tone — the preacher cares deeply for the people as they preach. Warm and substantive.",
        "delivery": "The congregation should feel spiritually fed, pastorally cared for, and theologically enriched.",
    },
    "problem-solution-tony-evans": {
        "name": "Problem–Solution (Practical & Direct)",
        "structure": "Clearly name the problem the congregation faces. Then unfold the biblical solution with precision. Move from diagnosis to prescription.",
        "style": "Direct and culturally aware. No hedging. Authoritative confidence that Scripture has the answer. Practical, actionable application.",
        "delivery": "The congregation should leave with a clear understanding of both their problem and God's specific solution — and know what to do about it.",
    },
    "homiletical-plot-terry-anderson": {
        "name": "Homiletical Plot (Journey to Celebration)",
        "structure": "Eugene Lowry's plot: Conflict → Complication → Clue → Gospel Turn → Resolution. Build toward a climactic Gospel moment.",
        "style": "Narrative movement with emotional escalation. Strong suspense through the middle. The Gospel turn arrives like a resolution the congregation has been aching for.",
        "delivery": "The congregation should experience a journey — conflict, tension, Gospel relief, and joyful celebration at the finish.",
    },
    "topical-hb-charles": {
        "name": "Topical (Organized & Focused)",
        "structure": "Multiple texts organized around a central topic or theme. Each point is anchored in Scripture. The topic shapes the structure; Scripture shapes each point.",
        "style": "Methodical and clearly organized. Disciplined movement between texts. No rambling. Clean outline with logical flow.",
        "delivery": "The congregation should leave with a comprehensive, biblically grounded understanding of the topic — organized well enough to teach others.",
    },
    "narrative-ralph-west": {
        "name": "Narrative (Reflective & Pastoral)",
        "structure": "Story-centered with pastoral pacing. The narrative arc is present but unhurried. Space for reflection within the story.",
        "style": "Pastoral warmth with reflective depth. Slower pacing to let truth settle. Relational tone — the preacher walks alongside the congregation through the story.",
        "delivery": "The congregation should feel the pastoral care embedded in the storytelling — moved emotionally, nourished spiritually.",
    },
    "big-idea-robert-smith": {
        "name": "Big Idea (Experiential & Imaginative)",
        "structure": "A single governing idea explored with depth and imagination. Uses the whole canon to illuminate the central truth.",
        "style": "Experiential and imaginative. Incarnational — the truth takes on flesh in language and story. Rich imagery. The idea is felt as much as understood.",
        "delivery": "The congregation should not only understand the big idea — they should have experienced it. It should live in them, not just in their notes.",
    },
}


@app.get("/api/sermon-styles")
async def get_sermon_styles():
    return [{"key": k, **{f: v[f] for f in ("name", "structure", "style", "delivery")}}
            for k, v in SERMON_STYLES.items()]


# ─── Export Helpers ───────────────────────────────────────────────────────────

def _inline_runs(text: str):
    """Split a markdown inline string into (content, style) tuples."""
    parts = re.split(r'(\*\*[^*\n]+\*\*|\*[^*\n]+\*|`[^`\n]+`)', text)
    result = []
    for part in parts:
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            result.append((part[2:-2], 'bold'))
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            result.append((part[1:-1], 'italic'))
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            result.append((part[1:-1], 'code'))
        elif part:
            result.append((part, 'normal'))
    return result


def markdown_to_docx(content: str, title: str) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches

    INK    = RGBColor(0x1B, 0x36, 0x5D)
    ACCENT = RGBColor(0x48, 0x7A, 0x7B)
    MUTED  = RGBColor(0x70, 0x73, 0x72)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    def add_heading(text: str, level: int):
        p = doc.add_heading('', level=level)
        p.clear()
        run = p.add_run(text.strip())
        run.font.bold = True
        run.font.color.rgb = INK if level <= 2 else ACCENT
        sizes = {1: 17, 2: 14, 3: 12, 4: 11}
        run.font.size = Pt(sizes.get(level, 11))
        return p

    def add_inline(text: str, style: str = 'Normal', indent: float = 0):
        p = doc.add_paragraph(style=style)
        if indent:
            p.paragraph_format.left_indent = Inches(indent)
        for chunk, kind in _inline_runs(text):
            run = p.add_run(chunk)
            run.font.color.rgb = INK
            if kind == 'bold':
                run.font.bold = True
            elif kind == 'italic':
                run.font.italic = True
            elif kind == 'code':
                run.font.name = 'Courier New'
                run.font.size = Pt(9.5)
        return p

    # Title block
    tp = doc.add_paragraph()
    tr = tp.add_run(title)
    tr.font.size = Pt(22)
    tr.font.bold = True
    tr.font.color.rgb = INK

    dp = doc.add_paragraph()
    dr = dp.add_run(f"Generated {datetime.now().strftime('%B %d, %Y')}  ·  Sermon Tools — IMB Innovation")
    dr.font.size = Pt(9)
    dr.font.color.rgb = MUTED

    doc.add_paragraph()

    for line in content.split('\n'):
        s = line.strip()
        if not s:
            doc.add_paragraph()
        elif s in ('---', '***', '___'):
            doc.add_paragraph('─' * 55)
        elif s.startswith('#### '):
            add_heading(s[5:], 4)
        elif s.startswith('### '):
            add_heading(s[4:], 3)
        elif s.startswith('## '):
            add_heading(s[3:], 2)
        elif s.startswith('# '):
            add_heading(s[2:], 1)
        elif s.startswith('- ') or s.startswith('* '):
            try:
                add_inline(s[2:], 'List Bullet')
            except Exception:
                add_inline(f'• {s[2:]}', indent=0.25)
        elif re.match(r'^\d+\.\s', s):
            try:
                add_inline(re.sub(r'^\d+\.\s', '', s), 'List Number')
            except Exception:
                add_inline(s, indent=0.25)
        elif s.startswith('> '):
            add_inline(s[2:], indent=0.3)
        else:
            add_inline(s)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def markdown_to_pdf(content: str, title: str) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    INK    = HexColor('#1B365D')
    ACCENT = HexColor('#487A7B')
    MUTED  = HexColor('#9a9c9a')
    BORDER = HexColor('#dde1e7')

    def S(name, **kw):
        return ParagraphStyle(name, **kw)

    title_s = S('T',  fontName='Helvetica-Bold',    fontSize=22, textColor=INK,    spaceAfter=4,  leading=26)
    meta_s  = S('Me', fontName='Helvetica',          fontSize=9,  textColor=MUTED,  spaceAfter=16)
    h1_s    = S('H1', fontName='Helvetica-Bold',    fontSize=17, textColor=INK,    spaceBefore=14, spaceAfter=6,  leading=22)
    h2_s    = S('H2', fontName='Helvetica-Bold',    fontSize=14, textColor=INK,    spaceBefore=12, spaceAfter=5,  leading=18)
    h3_s    = S('H3', fontName='Helvetica-Bold',    fontSize=12, textColor=ACCENT, spaceBefore=10, spaceAfter=4)
    h4_s    = S('H4', fontName='Helvetica-Bold',    fontSize=11, textColor=ACCENT, spaceBefore=8,  spaceAfter=3)
    body_s  = S('Bo', fontName='Helvetica',          fontSize=10.5, textColor=INK,  spaceAfter=5,  leading=16)
    bull_s  = S('Bu', fontName='Helvetica',          fontSize=10.5, textColor=INK,  spaceAfter=4,  leading=16, leftIndent=18)
    quot_s  = S('Qu', fontName='Helvetica-Oblique', fontSize=10,  textColor=HexColor('#444'),      spaceAfter=6,  leading=15, leftIndent=24, rightIndent=12)

    def fix(text: str) -> str:
        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
        text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
        text = re.sub(r'\*(.+?)\*',     r'<i>\1</i>', text)
        text = re.sub(r'`(.+?)`', r'<font face="Courier" size="9">\1</font>', text)
        return text

    story = [
        Paragraph(fix(title), title_s),
        Paragraph(f"Generated {datetime.now().strftime('%B %d, %Y')}  ·  Sermon Tools — IMB Innovation", meta_s),
        HRFlowable(width='100%', thickness=1, color=BORDER, spaceAfter=10),
    ]

    for line in content.split('\n'):
        s = line.strip()
        if not s:
            story.append(Spacer(1, 0.07 * inch))
        elif s in ('---', '***', '___'):
            story.append(HRFlowable(width='100%', thickness=0.5, color=BORDER, spaceBefore=4, spaceAfter=4))
        elif s.startswith('#### '):
            story.append(Paragraph(fix(s[5:]), h4_s))
        elif s.startswith('### '):
            story.append(Paragraph(fix(s[4:]), h3_s))
        elif s.startswith('## '):
            story.append(Paragraph(fix(s[3:]), h2_s))
        elif s.startswith('# '):
            story.append(Paragraph(fix(s[2:]), h1_s))
        elif s.startswith('- ') or s.startswith('* '):
            story.append(Paragraph(f'• {fix(s[2:])}', bull_s))
        elif re.match(r'^\d+\.\s', s):
            num, _, rest = s.partition('. ')
            story.append(Paragraph(f'{num}. {fix(rest)}', bull_s))
        elif s.startswith('> '):
            story.append(Paragraph(fix(s[2:]), quot_s))
        else:
            story.append(Paragraph(fix(s), body_s))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
        leftMargin=inch * 1.1, rightMargin=inch * 1.1,
        topMargin=inch * 1.0, bottomMargin=inch * 1.0)
    doc.build(story)
    return buf.getvalue()


# ─── Export Endpoint ──────────────────────────────────────────────────────────

class ExportRequest(BaseModel):
    content: str
    format: str
    title: str = "Sermon Tools"


@app.post("/api/export")
async def export_document(request: ExportRequest):
    if request.format not in ("docx", "pdf"):
        raise HTTPException(status_code=400, detail="Format must be 'docx' or 'pdf'")
    if not request.content.strip():
        raise HTTPException(status_code=400, detail="No content to export")

    safe_title = re.sub(r'[^\w\s-]', '', request.title)[:80].strip() or "Sermon Tools"

    if request.format == "docx":
        data = markdown_to_docx(request.content, request.title)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"{safe_title}.docx"
    else:
        data = markdown_to_pdf(request.content, request.title)
        media_type = "application/pdf"
        filename = f"{safe_title}.pdf"

    return Response(
        content=data,
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
